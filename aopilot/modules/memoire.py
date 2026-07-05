"""Module 5 — Générateur de mémoire technique.

Génération section par section (un appel API par grande section) en
suivant templates/memoire_template.md, avec règle absolue de ne jamais
inventer de données client, passe critique automatique, et export
.docx / .md.
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any, Callable

from .database import get_connection
from .dce import (
    MODELE_DCE,
    PRIX_INPUT_PAR_MTOK,
    PRIX_OUTPUT_PAR_MTOK,
    _appel_api,
    _client_anthropic,
    estimer_tokens,
)

CHEMIN_TEMPLATE: Path = (
    Path(__file__).resolve().parent.parent / "templates" / "memoire_template.md"
)
DOSSIER_OUTPUTS: Path = Path(__file__).resolve().parent.parent / "outputs"

# Règle absolue injectée dans chaque prompt de génération
REGLE_ANTI_INVENTION: str = (
    "RÈGLE ABSOLUE : ne JAMAIS inventer un chiffre, une certification, une "
    "référence client ou un nom propre. Si une information nécessaire manque "
    "dans la fiche entreprise, insère exactement le marqueur "
    "[DONNÉE CLIENT : description de l'information manquante] à sa place."
)

TOKENS_SORTIE_PAR_SECTION: int = 3000  # budget de sortie estimé par section


# ---------------------------------------------------------------------------
# Fiches entreprise
# ---------------------------------------------------------------------------

CHAMPS_FICHE: list[tuple[str, str]] = [
    ("raison_sociale", "Raison sociale"),
    ("siret", "SIRET"),
    ("effectif", "Effectif"),
    ("ca", "Chiffre d'affaires"),
    ("refs", "Références (marchés similaires)"),
    ("certifications", "Certifications (Qualipropre, MASE, ISO...)"),
    ("encadrement", "Encadrement (noms, rôles, qualifications)"),
    ("materiel", "Matériel et équipements"),
    ("produits", "Produits utilisés (écolabels...)"),
]


def sauvegarder_fiche(fiche: dict[str, str], fiche_id: int | None = None) -> int:
    """Crée ou met à jour une fiche entreprise. Retourne son id."""
    conn = get_connection()
    try:
        if fiche_id:
            conn.execute(
                """
                UPDATE fiches_entreprise SET raison_sociale=?, siret=?, effectif=?,
                    ca=?, refs=?, certifications=?, encadrement=?, materiel=?,
                    produits=?, date_maj=datetime('now')
                WHERE id=?
                """,
                (*[fiche.get(c, "") for c, _ in CHAMPS_FICHE], fiche_id),
            )
            conn.commit()
            return fiche_id
        curseur = conn.execute(
            """
            INSERT INTO fiches_entreprise
                (raison_sociale, siret, effectif, ca, refs, certifications,
                 encadrement, materiel, produits)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            tuple(fiche.get(c, "") for c, _ in CHAMPS_FICHE),
        )
        conn.commit()
        return int(curseur.lastrowid)
    finally:
        conn.close()


def lister_fiches() -> list[dict[str, Any]]:
    """Liste les fiches entreprises enregistrées."""
    conn = get_connection()
    try:
        lignes = conn.execute(
            "SELECT * FROM fiches_entreprise ORDER BY date_maj DESC"
        )
        return [dict(ligne) for ligne in lignes]
    finally:
        conn.close()


def fiche_en_texte(fiche: dict[str, Any]) -> str:
    """Formate une fiche entreprise en texte pour les prompts."""
    lignes = []
    for champ, libelle in CHAMPS_FICHE:
        valeur = fiche.get(champ) or "(non renseigné)"
        lignes.append(f"- {libelle} : {valeur}")
    return "\n".join(lignes)


# ---------------------------------------------------------------------------
# Template et sections
# ---------------------------------------------------------------------------


def charger_sections_template(chemin: Path | None = None) -> list[tuple[str, str]]:
    """Découpe le template markdown en sections (titre, instructions).

    Chaque section commence par '## Titre' ; le texte qui suit sert
    d'instructions de rédaction pour l'IA.
    """
    texte = (chemin or CHEMIN_TEMPLATE).read_text(encoding="utf-8")
    sections: list[tuple[str, str]] = []
    for bloc in re.split(r"^## ", texte, flags=re.MULTILINE)[1:]:
        lignes = bloc.strip().splitlines()
        titre = lignes[0].strip()
        instructions = "\n".join(lignes[1:]).strip()
        sections.append((titre, instructions))
    return sections


def estimer_cout_memoire(
    fiche: dict[str, Any], contexte_dce: str, nb_sections: int
) -> float:
    """Coût API estimé en dollars : nb_sections appels + 1 passe critique."""
    tokens_contexte = estimer_tokens(fiche_en_texte(fiche)) + estimer_tokens(
        contexte_dce
    )
    # Chaque appel de section renvoie le contexte + produit ~3000 tokens
    cout_sections = nb_sections * (
        tokens_contexte * PRIX_INPUT_PAR_MTOK / 1_000_000
        + TOKENS_SORTIE_PAR_SECTION * PRIX_OUTPUT_PAR_MTOK / 1_000_000
    )
    # Passe critique : relit le mémoire complet
    tokens_memoire = nb_sections * TOKENS_SORTIE_PAR_SECTION
    cout_critique = (
        tokens_memoire * PRIX_INPUT_PAR_MTOK / 1_000_000
        + 2000 * PRIX_OUTPUT_PAR_MTOK / 1_000_000
    )
    return cout_sections + cout_critique


# ---------------------------------------------------------------------------
# Génération
# ---------------------------------------------------------------------------


def generer_memoire(
    fiche: dict[str, Any],
    contexte_dce: str,
    titre_marche: str,
    rappel_progression: Callable[[int, int, str], None] | None = None,
) -> str:
    """Génère le mémoire complet, section par section.

    rappel_progression(i, total, titre_section) permet d'alimenter la barre
    de progression Streamlit.
    """
    client = _client_anthropic()
    sections = charger_sections_template()
    texte_fiche = fiche_en_texte(fiche)

    parties: list[str] = [f"# Mémoire technique — {titre_marche}"]
    for i, (titre, instructions) in enumerate(sections, start=1):
        if rappel_progression:
            rappel_progression(i, len(sections), titre)
        systeme = (
            "Tu es un rédacteur professionnel de mémoires techniques d'appels "
            "d'offres publics français, spécialisé TPE/PME (nettoyage, BTP second "
            "œuvre, espaces verts, sécurité). Tu rédiges en français, de façon "
            "concrète et orientée critères de notation.\n" + REGLE_ANTI_INVENTION
        )
        contenu = (
            f"MARCHÉ : {titre_marche}\n\n"
            f"FICHE ENTREPRISE CANDIDATE :\n{texte_fiche}\n\n"
            f"CONTEXTE DU DCE (analyse) :\n{contexte_dce or '(aucune analyse fournie)'}\n\n"
            f"SECTION À RÉDIGER : « {titre} »\n"
            f"CONSIGNES DE LA SECTION :\n{instructions}\n\n"
            "Rédige uniquement le contenu de cette section, en markdown, sans "
            "répéter le titre de section (il sera ajouté automatiquement)."
        )
        texte_section = _appel_api(client, systeme, contenu, max_tokens=4000)
        parties.append(f"## {titre}\n\n{texte_section.strip()}")

    return "\n\n".join(parties)


def critiquer_memoire(memoire_md: str, contexte_rc: str) -> str:
    """Passe critique : note le mémoire selon les critères du RC et liste les faiblesses."""
    client = _client_anthropic()
    systeme = (
        "Tu es un acheteur public expérimenté qui note des mémoires techniques. "
        "Tu es exigeant et concret."
    )
    contenu = (
        f"CRITÈRES DE NOTATION (issus du règlement de consultation) :\n"
        f"{contexte_rc or '(non fournis — utilise les critères usuels : valeur technique, moyens, méthodologie, RSE)'}\n\n"
        f"MÉMOIRE À ÉVALUER :\n{memoire_md}\n\n"
        "Produis :\n"
        "1) une note estimée par critère (avec justification),\n"
        "2) la liste des faiblesses classées par impact sur la note,\n"
        "3) la liste des marqueurs [DONNÉE CLIENT : ...] restants à compléter,\n"
        "4) 5 améliorations prioritaires."
    )
    return _appel_api(client, systeme, contenu, max_tokens=4000)


# ---------------------------------------------------------------------------
# Sauvegarde et exports
# ---------------------------------------------------------------------------


def sauvegarder_memoire(
    titre: str,
    contenu_md: str,
    critique: str = "",
    marche_idweb: str | None = None,
    fiche_id: int | None = None,
    statut: str = "en cours",
) -> int:
    """Enregistre le mémoire en base. Retourne son id."""
    conn = get_connection()
    try:
        curseur = conn.execute(
            """
            INSERT INTO memoires (marche_idweb, fiche_id, titre, contenu_md, critique, statut)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (marche_idweb, fiche_id, titre, contenu_md, critique, statut),
        )
        conn.commit()
        return int(curseur.lastrowid)
    finally:
        conn.close()


def mettre_a_jour_memoire(memoire_id: int, contenu_md: str) -> None:
    """Sauvegarde le contenu édité d'un mémoire existant."""
    conn = get_connection()
    try:
        conn.execute(
            "UPDATE memoires SET contenu_md = ? WHERE id = ?", (contenu_md, memoire_id)
        )
        conn.commit()
    finally:
        conn.close()


def supprimer_memoire(memoire_id: int) -> None:
    """Supprime un mémoire de la base."""
    conn = get_connection()
    try:
        conn.execute("DELETE FROM memoires WHERE id = ?", (memoire_id,))
        conn.commit()
    finally:
        conn.close()


def lister_memoires() -> list[dict[str, Any]]:
    """Liste les mémoires enregistrés (le plus récent d'abord)."""
    conn = get_connection()
    try:
        lignes = conn.execute("SELECT * FROM memoires ORDER BY date_creation DESC")
        return [dict(ligne) for ligne in lignes]
    finally:
        conn.close()


def exporter_md(contenu_md: str, nom_fichier: str) -> Path:
    """Écrit le mémoire en .md dans outputs/ et retourne le chemin."""
    DOSSIER_OUTPUTS.mkdir(exist_ok=True)
    chemin = DOSSIER_OUTPUTS / f"{nom_fichier}.md"
    chemin.write_text(contenu_md, encoding="utf-8")
    return chemin


def _ajouter_paragraphe_md(document: Any, ligne: str) -> None:
    """Ajoute une ligne markdown simple (gras **x**) dans un paragraphe docx."""
    paragraphe = document.add_paragraph()
    # Découpe sur les segments en gras pour préserver l'emphase
    for i, segment in enumerate(re.split(r"\*\*(.+?)\*\*", ligne)):
        if not segment:
            continue
        run = paragraphe.add_run(segment)
        run.bold = i % 2 == 1  # les indices impairs sont entre **...**


def exporter_docx(
    contenu_md: str, nom_fichier: str, entreprise: str = "", marche: str = ""
) -> Path:
    """Convertit le mémoire markdown en .docx : page de garde, sommaire,
    styles de titres et tableaux. Retourne le chemin du fichier."""
    from docx import Document  # import local : python-docx requis pour l'export
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    DOSSIER_OUTPUTS.mkdir(exist_ok=True)
    document = Document()

    # --- Page de garde ---
    titre = document.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("MÉMOIRE TECHNIQUE")
    run.bold = True
    run.font.size = Pt(28)
    for texte in (marche, entreprise, date.today().strftime("%d/%m/%Y")):
        if texte:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(texte).font.size = Pt(14)
    document.add_page_break()

    # --- Sommaire (champ TOC Word, mis à jour à l'ouverture du document) ---
    document.add_heading("Sommaire", level=1)
    paragraphe_toc = document.add_paragraph()
    fld = paragraphe_toc.add_run()._r
    fld_char_begin = fld.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instr = fld.makeelement(qn("w:instrText"), {})
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_char_end = fld.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    for element in (fld_char_begin, instr, fld_char_end):
        fld.append(element)
    document.add_page_break()

    # --- Corps : conversion markdown ligne à ligne ---
    lignes = contenu_md.splitlines()
    i = 0
    while i < len(lignes):
        ligne = lignes[i].rstrip()
        if ligne.startswith("|"):
            # Bloc tableau markdown : collecte des lignes consécutives
            bloc: list[str] = []
            while i < len(lignes) and lignes[i].strip().startswith("|"):
                bloc.append(lignes[i].strip())
                i += 1
            _ajouter_tableau_docx(document, bloc)
            continue
        if ligne.startswith("### "):
            document.add_heading(ligne[4:], level=3)
        elif ligne.startswith("## "):
            document.add_heading(ligne[3:], level=1)
        elif ligne.startswith("# "):
            pass  # le titre principal est déjà sur la page de garde
        elif ligne.startswith(("- ", "* ")):
            document.add_paragraph(ligne[2:], style="List Bullet")
        elif re.match(r"^\d+[.)] ", ligne):
            document.add_paragraph(re.sub(r"^\d+[.)] ", "", ligne), style="List Number")
        elif ligne:
            _ajouter_paragraphe_md(document, ligne)
        i += 1

    chemin = DOSSIER_OUTPUTS / f"{nom_fichier}.docx"
    document.save(str(chemin))
    return chemin


def _ajouter_tableau_docx(document: Any, lignes_md: list[str]) -> None:
    """Convertit un bloc de tableau markdown en tableau Word."""
    # Ignore la ligne séparatrice |---|---|
    lignes_utiles = [
        l for l in lignes_md if not re.match(r"^\|[\s:|-]+\|$", l.replace(" ", ""))
    ]
    if not lignes_utiles:
        return
    cellules = [
        [c.strip() for c in ligne.strip("|").split("|")] for ligne in lignes_utiles
    ]
    nb_colonnes = max(len(l) for l in cellules)
    tableau = document.add_table(rows=len(cellules), cols=nb_colonnes)
    tableau.style = "Table Grid"
    for r, ligne in enumerate(cellules):
        for c, valeur in enumerate(ligne):
            cellule = tableau.cell(r, c)
            cellule.text = valeur.replace("**", "")
            if r == 0:  # en-tête en gras
                for paragraphe in cellule.paragraphs:
                    for run in paragraphe.runs:
                        run.bold = True
