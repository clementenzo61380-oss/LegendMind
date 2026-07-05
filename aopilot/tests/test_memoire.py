"""Tests du module mémoire : découpage du template et exports (sans appel API)."""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from modules.memoire import (
    charger_sections_template,
    estimer_cout_memoire,
    exporter_docx,
    exporter_md,
    fiche_en_texte,
)


def test_template_par_defaut_a_8_sections() -> None:
    sections = charger_sections_template()
    titres = [t for t, _ in sections]
    assert len(sections) == 8
    assert titres[0] == "Présentation de l'entreprise"
    assert titres[-1] == "Hygiène et sécurité"
    # Chaque section a des consignes non vides pour guider l'IA
    assert all(instructions for _, instructions in sections)


def test_template_personnalise(tmp_path: Path) -> None:
    """Un template fourni par l'utilisateur est découpé sur les '## '."""
    chemin = tmp_path / "perso.md"
    chemin.write_text(
        "# Mon mémoire\n\n## Section A\nConsigne A\n\n## Section B\nConsigne B\n",
        encoding="utf-8",
    )
    sections = charger_sections_template(chemin)
    assert sections == [("Section A", "Consigne A"), ("Section B", "Consigne B")]


def test_fiche_en_texte_signale_champs_vides() -> None:
    texte = fiche_en_texte({"raison_sociale": "Test SARL"})
    assert "Raison sociale : Test SARL" in texte
    assert "(non renseigné)" in texte  # les champs vides restent visibles


def test_estimation_cout_positive() -> None:
    cout = estimer_cout_memoire({"raison_sociale": "X"}, "contexte" * 200, 8)
    assert 0 < cout < 5  # ordre de grandeur raisonnable en dollars


def test_export_md_et_docx(tmp_path: Path, monkeypatch) -> None:
    """Les deux exports produisent des fichiers non vides ; le docx gère les tableaux."""
    import modules.memoire as memoire_mod

    monkeypatch.setattr(memoire_mod, "DOSSIER_OUTPUTS", tmp_path)
    md = (
        "# Mémoire technique — Test\n\n"
        "## Moyens humains\n\nUne équipe **dédiée** de 5 agents.\n\n"
        "| Poste | Effectif |\n|---|---|\n| Agent | 4 |\n| Chef | 1 |\n\n"
        "- EPI fournis\n1. Étape un\n"
    )
    chemin_md = memoire_mod.exporter_md(md, "test")
    assert chemin_md.read_text(encoding="utf-8") == md

    chemin_docx = memoire_mod.exporter_docx(md, "test", "Test SARL", "Marché test")
    assert chemin_docx.exists() and chemin_docx.stat().st_size > 1000

    # Vérifie que le tableau markdown est bien devenu un tableau Word
    from docx import Document

    document = Document(str(chemin_docx))
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "Poste"
    assert document.tables[0].cell(2, 1).text == "1"
