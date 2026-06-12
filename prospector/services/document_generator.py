"""
Génération de contrats d'apporteur d'affaires.
Format : .docx (python-docx) et .pdf (reportlab)
⚠️ AVERTISSEMENT LÉGAL : Ces modèles sont fournis à titre indicatif.
   Faites valider tout contrat par un professionnel du droit.
"""
import os
from pathlib import Path
from datetime import datetime
from typing import Optional

OUTPUT_DIR = Path("generated_docs")

APPORTEUR_NOM = os.getenv("APPORTEUR_NOM", "Prénom NOM")
APPORTEUR_SIREN = os.getenv("APPORTEUR_SIREN", "000 000 000")
APPORTEUR_ADRESSE = os.getenv("APPORTEUR_ADRESSE", "Adresse, 35000 Rennes")
APPORTEUR_TEL = os.getenv("APPORTEUR_TELEPHONE", "06 XX XX XX XX")
APPORTEUR_EMAIL = os.getenv("APPORTEUR_EMAIL", "contact@exemple.fr")
COMMISSION_DEFAUT = float(os.getenv("COMMISSION_DEFAUT", "5.0"))

DISCLAIMER = (
    "⚠️ MODÈLE À FAIRE VALIDER PAR UN PROFESSIONNEL DU DROIT — "
    "Ce document est fourni à titre indicatif uniquement."
)


def _build_contract_text(artisan: dict, taux_commission: float) -> dict:
    """Construit le texte du contrat avec les variables remplies."""
    today = datetime.now().strftime("%d %B %Y")
    artisan_nom = artisan.get("nom", "")
    artisan_siren = artisan.get("siren", "")
    artisan_adresse = f"{artisan.get('ville', 'Rennes')}"
    artisan_metier = artisan.get("metier", "artisan du bâtiment")

    title = "CONTRAT D'APPORTEUR D'AFFAIRES — BTP / RÉNOVATION"

    sections = [
        {
            "titre": "ENTRE LES SOUSSIGNÉS",
            "contenu": f"""L'apporteur d'affaires :
Nom / Raison sociale : {APPORTEUR_NOM}
SIREN : {APPORTEUR_SIREN}
Adresse : {APPORTEUR_ADRESSE}
Téléphone : {APPORTEUR_TEL}
Email : {APPORTEUR_EMAIL}
Ci-après dénommé « L'Apporteur »

ET

L'artisan partenaire :
Nom / Raison sociale : {artisan_nom}
SIREN : {artisan_siren}
Adresse : {artisan_adresse}
Activité : {artisan_metier}
Ci-après dénommé « L'Artisan »

IL A ÉTÉ CONVENU CE QUI SUIT :"""
        },
        {
            "titre": "ARTICLE 1 — OBJET DU CONTRAT",
            "contenu": """Le présent contrat a pour objet de définir les conditions dans lesquelles l'Apporteur
met en relation l'Artisan avec des prospects ayant exprimé un besoin en travaux de
bâtiment et/ou de rénovation énergétique (ci-après « Leads »).

L'Apporteur agit en qualité d'intermédiaire indépendant, sans mandat de représentation.
Il ne se substitue pas à l'Artisan dans la réalisation des prestations techniques."""
        },
        {
            "titre": "ARTICLE 2 — DÉFINITION DU LEAD QUALIFIÉ",
            "contenu": """Un « Lead qualifié » au sens du présent contrat est tout contact transmis par
l'Apporteur qui réunit les conditions suivantes :
a) La personne a exprimé un besoin explicite en travaux relevant de l'activité de l'Artisan ;
b) Les coordonnées (nom, téléphone et/ou email) ont été vérifiées et sont joignables ;
c) Le projet est localisé dans la zone d'intervention de l'Artisan.

La qualification est attestée par l'email de transmission horodaté envoyé par l'Apporteur."""
        },
        {
            "titre": "ARTICLE 3 — RÉMUNÉRATION DE L'APPORTEUR",
            "contenu": f"""3.1 — Taux de commission
En contrepartie de chaque mise en relation ayant abouti à la signature d'un devis accepté
ET au versement d'un acompte par le client, l'Artisan versera à l'Apporteur une commission de :

    {taux_commission:.1f}% ({"".join(["zéro virgule cinq" if taux_commission == 0.5 else str(int(taux_commission))]) + " virgule " + str(int((taux_commission % 1) * 10)) if taux_commission % 1 != 0 else str(int(taux_commission))} pour cent)

du montant HT du devis signé.

3.2 — Fait générateur
La commission est due à la date à laquelle les deux conditions suivantes sont réunies :
- le devis a été signé par le client ;
- l'acompte contractuel a été versé par le client à l'Artisan.

3.3 — Délai de paiement
L'Artisan s'engage à régler la commission dans un délai de 30 (trente) jours calendaires
à compter du fait générateur défini à l'article 3.2.

3.4 — Modalités
Paiement par virement bancaire sur le RIB communiqué par l'Apporteur.
Toute facture impayée au-delà de 45 jours portera intérêts de retard au taux légal majoré
de 3 points, sans mise en demeure préalable."""
        },
        {
            "titre": "ARTICLE 4 — OBLIGATION D'INFORMATION DE L'ARTISAN",
            "contenu": """L'Artisan s'engage à informer l'Apporteur, dans un délai de 7 (sept) jours :
a) de la prise de contact avec le prospect ;
b) de l'émission du devis (en communiquant le montant HT) ;
c) de la signature du devis et du versement de l'acompte ;
d) de tout abandon ou refus du projet par le prospect.

L'absence d'information dans ce délai pourra entraîner la suspension des transmissions
de leads jusqu'à régularisation."""
        },
        {
            "titre": "ARTICLE 5 — CLAUSE DE NON-CONTOURNEMENT",
            "contenu": """L'Artisan s'interdit formellement de traiter directement ou indirectement avec tout
prospect transmis par l'Apporteur, sans que la commission prévue à l'article 3 soit réglée,
et ce pendant une durée de 24 (vingt-quatre) mois à compter de la date de transmission.

Cette interdiction s'applique à tout contrat, devis ou prestation, même réalisé sous
une autre dénomination sociale ou par l'intermédiaire d'un tiers.

En cas de violation, l'Artisan sera redevable d'une indemnité forfaitaire égale à
3 fois le montant de la commission qui aurait été due."""
        },
        {
            "titre": "ARTICLE 6 — DURÉE ET RÉSILIATION",
            "contenu": """Le présent contrat est conclu pour une durée indéterminée à compter de sa signature.
Il peut être résilié par l'une ou l'autre des parties avec un préavis de 30 jours,
par lettre recommandée avec accusé de réception.

La résiliation ne remet pas en cause les droits à commission acquis sur les leads
déjà transmis avant la date d'effet de la résiliation."""
        },
        {
            "titre": "ARTICLE 7 — CONFIDENTIALITÉ",
            "contenu": """Chaque partie s'engage à garder confidentielles les informations relatives aux
clients, prospects et partenaires de l'autre partie, pendant toute la durée du contrat
et pour une durée de 3 ans après son terme."""
        },
        {
            "titre": "ARTICLE 8 — LOI APPLICABLE — LITIGE",
            "contenu": """Le présent contrat est soumis au droit français.
En cas de litige, les parties s'engagent à rechercher une solution amiable.
À défaut, le Tribunal de Commerce de Rennes sera seul compétent."""
        },
        {
            "titre": "SIGNATURES",
            "contenu": f"""Fait à Rennes, le {today}, en deux exemplaires originaux.

Pour L'Apporteur :                        Pour L'Artisan :
{APPORTEUR_NOM}                           {artisan_nom}

Signature :                               Signature :


___________________________               ___________________________
(Faire précéder de la mention             (Faire précéder de la mention
 « Lu et approuvé »)                       « Lu et approuvé »)"""
        },
    ]

    return {
        "title": title,
        "disclaimer": DISCLAIMER,
        "date": today,
        "sections": sections,
        "artisan": artisan,
        "apporteur_nom": APPORTEUR_NOM,
        "taux_commission": taux_commission,
    }


def generate_docx(artisan: dict, taux_commission: Optional[float] = None) -> Path:
    """Génère un contrat .docx et retourne le chemin du fichier."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx non installé. Lancez : pip install python-docx")

    if taux_commission is None:
        taux_commission = artisan.get("taux_commission", COMMISSION_DEFAUT)

    contract = _build_contract_text(artisan, taux_commission)
    OUTPUT_DIR.mkdir(exist_ok=True)

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Disclaimer en rouge en haut
    p = doc.add_paragraph()
    run = p.add_run(contract["disclaimer"])
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Titre
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(contract["title"])
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Sections
    for section in contract["sections"]:
        # Titre section
        h = doc.add_paragraph()
        h_run = h.add_run(section["titre"])
        h_run.bold = True
        h_run.font.size = Pt(12)
        h_run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

        # Contenu — formatage sur les runs, pas sur le style partagé
        p = doc.add_paragraph(section["contenu"])
        for run in p.runs:
            run.font.size = Pt(11)

        doc.add_paragraph()

    # Nom du fichier
    artisan_slug = "".join(c for c in artisan.get("nom", "artisan") if c.isalnum() or c in " -")[:30].strip()
    filename = f"contrat_{artisan_slug.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = OUTPUT_DIR / filename

    doc.save(str(filepath))
    return filepath


def generate_pdf(artisan: dict, taux_commission: Optional[float] = None) -> Path:
    """Génère un contrat .pdf et retourne le chemin du fichier."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    except ImportError:
        raise RuntimeError("reportlab non installé. Lancez : pip install reportlab")

    if taux_commission is None:
        taux_commission = artisan.get("taux_commission", COMMISSION_DEFAUT)

    contract = _build_contract_text(artisan, taux_commission)
    OUTPUT_DIR.mkdir(exist_ok=True)

    artisan_slug = "".join(c for c in artisan.get("nom", "artisan") if c.isalnum() or c in " -")[:30].strip()
    filename = f"contrat_{artisan_slug.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
    filepath = OUTPUT_DIR / filename

    doc = SimpleDocTemplate(
        str(filepath),
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    style_disclaimer = ParagraphStyle(
        "disclaimer", parent=styles["Normal"],
        fontSize=9, textColor=colors.red, alignment=TA_CENTER, spaceAfter=12
    )
    style_title = ParagraphStyle(
        "title", parent=styles["Title"],
        fontSize=16, textColor=colors.HexColor("#1e3a5f"),
        alignment=TA_CENTER, spaceAfter=20
    )
    style_section_title = ParagraphStyle(
        "section_title", parent=styles["Heading2"],
        fontSize=12, textColor=colors.HexColor("#1e3a5f"),
        spaceBefore=14, spaceAfter=6, fontName="Helvetica-Bold"
    )
    style_body = ParagraphStyle(
        "body", parent=styles["Normal"],
        fontSize=10, leading=16, alignment=TA_JUSTIFY, spaceAfter=8
    )

    story = []

    story.append(Paragraph(contract["disclaimer"], style_disclaimer))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(contract["title"], style_title))
    story.append(Spacer(1, 0.5 * cm))

    for section in contract["sections"]:
        story.append(Paragraph(section["titre"], style_section_title))
        # Échapper les entités XML puis convertir les sauts de ligne en <br/>
        import html as _html
        contenu_safe = _html.escape(section["contenu"]).replace("\n", "<br/>")
        story.append(Paragraph(contenu_safe, style_body))

    doc.build(story)
    return filepath
